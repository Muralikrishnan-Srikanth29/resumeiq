"""
Raw text extraction from uploaded files. This is step 1 of the
mandatory pipeline: PDF/DOCX → plain text → (later) structured JSON.
No file bytes or PDF content ever reach the AI layer — only the text
extracted here, and only after it's been structured.
"""
import io

import fitz  # PyMuPDF
from docx import Document

from app.core.exceptions import ParsingError, UnsupportedFileTypeError


def extract_text_from_pdf(file_bytes: bytes) -> str:
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        pages = [page.get_text("text") for page in doc]
        doc.close()
        text = "\n".join(pages).strip()
        if not text:
            raise ParsingError(
                "No extractable text found. This PDF may be a scanned image — "
                "OCR is not currently supported."
            )
        return text
    except ParsingError:
        raise
    except Exception as exc:
        raise ParsingError(f"Failed to read PDF: {exc}") from exc


def extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        doc = Document(io.BytesIO(file_bytes))
        parts: list[str] = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)
        text = "\n".join(parts).strip()
        if not text:
            raise ParsingError("No extractable text found in DOCX file.")
        return text
    except ParsingError:
        raise
    except Exception as exc:
        raise ParsingError(f"Failed to read DOCX: {exc}") from exc


def extract_text(file_bytes: bytes, filename: str) -> tuple[str, str]:
    """Returns (extracted_text, file_type)."""
    lower = filename.lower()
    if lower.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes), "pdf"
    if lower.endswith(".docx"):
        return extract_text_from_docx(file_bytes), "docx"
    raise UnsupportedFileTypeError()
